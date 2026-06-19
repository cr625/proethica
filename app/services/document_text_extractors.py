"""
Document text extraction (PDF / DOCX / TXT / HTML / URL).

Pure text-extraction functions with no embedding-model or database state. Split
out of ``EmbeddingService`` (where they lived as private methods) so the
embedding service file is not also a document parser. ``EmbeddingService`` keeps
thin ``_extract_text`` / ``_extract_from_url`` methods that delegate here, so the
existing instance-method call sites (case-creation document pipeline, task_queue)
are unchanged.

Heavy parsers (PyPDF2, python-docx, BeautifulSoup) are imported lazily inside
each function, matching the original behavior: a missing parser raises a clear
ImportError only when that format is actually requested.
"""

import logging

import requests

logger = logging.getLogger(__name__)


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract text from a file based on its type.

    Args:
        file_path: Path to the file
        file_type: Type of the file (pdf, docx, txt, html, url)

    Returns:
        Extracted text content
    """
    file_type = file_type.lower() if file_type else ''

    try:
        # Handle different file types
        if file_type == 'pdf':
            return extract_from_pdf(file_path)
        elif file_type == 'docx':
            return extract_from_docx(file_path)
        elif file_type in ['txt', 'text']:
            return extract_from_txt(file_path)
        elif file_type in ['html', 'htm']:
            return extract_from_html(file_path)
        elif file_type == 'url':
            return extract_from_url(file_path)  # In this case, file_path would be the URL
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    except Exception as e:
        # Log the error and re-raise
        logger.error(f"Error extracting text from {file_path} ({file_type}): {str(e)}")
        raise


def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        from PyPDF2 import PdfReader

        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = ""

            # Extract text from each page
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

            return text.strip()
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF text extraction. Install it with 'pip install PyPDF2'")


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        import docx

        doc = docx.Document(file_path)
        text = ""

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"

        return text.strip()
    except ImportError:
        raise ImportError("python-docx is required for DOCX text extraction. Install it with 'pip install python-docx'")


def extract_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def extract_from_html(file_path: str) -> str:
    """Extract text from HTML file."""
    try:
        from bs4 import BeautifulSoup

        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()

            # Get text
            text = soup.get_text()

            # Break into lines and remove leading and trailing space on each
            lines = (line.strip() for line in text.splitlines())

            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))

            # Drop blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text
    except ImportError:
        raise ImportError("BeautifulSoup4 is required for HTML text extraction. Install it with 'pip install beautifulsoup4'")


def extract_from_url(url: str) -> str:
    """Extract text from a URL, preserving structure and numbering."""
    try:
        from bs4 import BeautifulSoup
        import re

        # Fetch URL content
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        # Parse with BeautifulSoup
        soup = BeautifulSoup(response.content, 'html.parser')

        # Remove script, style elements, and hidden elements
        for element in soup(["script", "style", "head", "meta", "noscript"]):
            element.extract()

        # Process and extract text while preserving structure
        result_text = []

        # Handle headings
        for heading in soup.find_all(re.compile('^h[1-6]$')):
            heading_text = heading.get_text().strip()
            if heading_text:
                result_text.append(f"\n{'#' * int(heading.name[1])} {heading_text}\n")

        # Handle lists
        for list_element in soup.find_all(['ul', 'ol']):
            for i, item in enumerate(list_element.find_all('li'), 1):
                item_text = item.get_text().strip()
                if list_element.name == 'ul':
                    result_text.append(f"• {item_text}")
                else:
                    result_text.append(f"{i}. {item_text}")

        # Handle paragraphs with potential numbering
        for paragraph in soup.find_all(['p', 'div']):
            # Skip if it's a child of lists or headings we've already processed
            if paragraph.find_parent(['ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                continue

            text = paragraph.get_text().strip()
            if text:
                # Preserve paragraph numbering if present
                numbered_match = re.match(r'^(\d+\.?\d*|\(\d+\)|\w+\.)\s+(.+)$', text)
                if numbered_match:
                    number, content = numbered_match.groups()
                    result_text.append(f"{number} {content}")
                else:
                    result_text.append(text)

        # Special handling for tables
        for table in soup.find_all('table'):
            result_text.append("\n--- TABLE ---\n")
            for row in table.find_all('tr'):
                cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                result_text.append(" | ".join(cells))
            result_text.append("--- END TABLE ---\n")

        # Combine everything
        combined_text = "\n\n".join(result_text)

        # Clean up extra whitespace while preserving structure
        cleaned_lines = []
        for line in combined_text.split('\n'):
            line = re.sub(r'\s+', ' ', line).strip()
            if line:
                cleaned_lines.append(line)

        final_text = '\n'.join(cleaned_lines)

        logger.debug(f"Extracted text from URL: {url}")
        logger.debug(f"Text length: {len(final_text)} characters")

        return final_text
    except ImportError:
        raise ImportError("BeautifulSoup4 is required for URL text extraction. Install it with 'pip install beautifulsoup4'")
    except requests.RequestException as e:
        raise ValueError(f"Error fetching URL {url}: {str(e)}")
