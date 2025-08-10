import os
import numpy as np
from typing import List, Dict, Any, Union, Optional
import requests
import json
from sqlalchemy import text
import io
import logging

# Set up logging
logger = logging.getLogger(__name__)

class EmbeddingService:
    """
    Service for generating and managing embeddings for RDF triples.
    Supports using local or remote embedding models with configurable provider priority.
    """
    
    def __init__(self, model_name=None, embedding_dimension=None):
        """
        Initialize the embedding service.
        
        Args:
            model_name: The name of the local embedding model to use (defaults to env var or 'all-MiniLM-L6-v2')
            embedding_dimension: The dimension of the embedding vectors (determined by model)
        """
        # Configuration from environment or default values
        self.model_name = model_name or os.environ.get("LOCAL_EMBEDDING_MODEL", "all-MiniLM-L6-v2")
        
        # Provider priority (configurable through environment)
        self.provider_priority = os.environ.get(
            "EMBEDDING_PROVIDER_PRIORITY",
            # Default keeps current behavior; set to "openai,gemini,claude,local" to prefer hosted
            "local,claude,openai"
        ).lower().split(',')
        
        # Model dimensions (these will be used if embeddings need to be generated from scratch)
        self.dimensions = {
            "local": embedding_dimension or 384,   # all-MiniLM-L6-v2
            "claude": 1024,                       # Claude embeddings
            "openai": 1536,                      # text-embedding-ada-002
            "gemini": 768                        # text-embedding-004
        }
        
        # Default dimension based on first provider in priority
        for provider in self.provider_priority:
            if provider in self.dimensions:
                self.embedding_dimension = self.dimensions[provider]
                break
        else:
            self.embedding_dimension = embedding_dimension or 384  # Fallback
        
        # Provider setup and validation
        self.providers = {}
        self._setup_providers()
        
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from a file based on its type.
        
        Args:
            file_path: Path to the file
            file_type: Type of the file (pdf, docx, txt, html, url)
            
        Returns:
            Extracted text content
        """
        file_type = file_type.lower() if file_type else ''
        
        try:
            # Handle different file types
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type == 'docx':
                return self._extract_from_docx(file_path)
            elif file_type in ['txt', 'text']:
                return self._extract_from_txt(file_path)
            elif file_type in ['html', 'htm']:
                return self._extract_from_html(file_path)
            elif file_type == 'url':
                return self._extract_from_url(file_path)  # In this case, file_path would be the URL
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            # Log the error and re-raise
            print(f"Error extracting text from {file_path} ({file_type}): {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            from PyPDF2 import PdfReader
            
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                text = ""
                
                # Extract text from each page
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                return text.strip()
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF text extraction. Install it with 'pip install PyPDF2'")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except ImportError:
            raise ImportError("python-docx is required for DOCX text extraction. Install it with 'pip install python-docx'")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                # Get text
                text = soup.get_text()
                
                # Break into lines and remove leading and trailing space on each
                lines = (line.strip() for line in text.splitlines())
                
                # Break multi-headlines into a line each
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                
                # Drop blank lines
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
        except ImportError:
            raise ImportError("BeautifulSoup4 is required for HTML text extraction. Install it with 'pip install beautifulsoup4'")
    
    def _extract_from_url(self, url: str) -> str:
        """Extract text from a URL, preserving structure and numbering."""
        try:
            from bs4 import BeautifulSoup
            import re
            
            # Fetch URL content
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
            }
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script, style elements, and hidden elements
            for element in soup(["script", "style", "head", "meta", "noscript"]):
                element.extract()
                
            # Process and extract text while preserving structure
            result_text = []
            
            # Handle headings
            for heading in soup.find_all(re.compile('^h[1-6]$')):
                heading_text = heading.get_text().strip()
                if heading_text:
                    result_text.append(f"\n{'#' * int(heading.name[1])} {heading_text}\n")
            
            # Handle lists
            for list_element in soup.find_all(['ul', 'ol']):
                for i, item in enumerate(list_element.find_all('li'), 1):
                    item_text = item.get_text().strip()
                    if list_element.name == 'ul':
                        result_text.append(f"• {item_text}")
                    else:
                        result_text.append(f"{i}. {item_text}")
            
            # Handle paragraphs with potential numbering
            for paragraph in soup.find_all(['p', 'div']):
                # Skip if it's a child of lists or headings we've already processed
                if paragraph.find_parent(['ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    continue
                    
                text = paragraph.get_text().strip()
                if text:
                    # Preserve paragraph numbering if present
                    numbered_match = re.match(r'^(\d+\.?\d*|\(\d+\)|\w+\.)\s+(.+)$', text)
                    if numbered_match:
                        number, content = numbered_match.groups()
                        result_text.append(f"{number} {content}")
                    else:
                        result_text.append(text)
            
            # Special handling for tables
            for table in soup.find_all('table'):
                result_text.append("\n--- TABLE ---\n")
                for row in table.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    result_text.append(" | ".join(cells))
                result_text.append("--- END TABLE ---\n")
            
            # Combine everything
            combined_text = "\n\n".join(result_text)
            
            # Clean up extra whitespace while preserving structure
            cleaned_lines = []
            for line in combined_text.split('\n'):
                line = re.sub(r'\s+', ' ', line).strip()
                if line:
                    cleaned_lines.append(line)
            
            final_text = '\n'.join(cleaned_lines)
            
            # Debug logging
            print(f"Extracted text from URL: {url}")
            print(f"Text length: {len(final_text)} characters")
            
            return final_text
        except ImportError:
            raise ImportError("BeautifulSoup4 is required for URL text extraction. Install it with 'pip install beautifulsoup4'")
        except requests.RequestException as e:
            raise ValueError(f"Error fetching URL {url}: {str(e)}")
    
    def _setup_providers(self):
        """Initialize and validate all configured providers."""
        # Local model setup
        if "local" in self.provider_priority:
            try:
                from sentence_transformers import SentenceTransformer
                import torch
                
                # Configure offline mode to avoid HuggingFace Hub requests
                os.environ["HF_HUB_OFFLINE"] = "1"
                os.environ["TRANSFORMERS_OFFLINE"] = "1"
                # Device selection with CPU fallback option
                requested_device = os.environ.get("EMBEDDINGS_DEVICE", "auto").lower()
                device = "cpu" if requested_device == "cpu" else ("cuda" if torch.cuda.is_available() and requested_device != "cpu" else "cpu")
                
                self.providers["local"] = {
                    "model": SentenceTransformer(self.model_name, local_files_only=True, device=device),
                    "available": True,
                    "dimension": self.dimensions["local"],
                    "device": device
                }
                print(f"Local embedding provider ready: {self.model_name} (device={device})")
            except Exception as e:
                print(f"Local embedding provider unavailable: {str(e)}")
                self.providers["local"] = {"available": False}
        
        # Claude API setup
        if "claude" in self.provider_priority:
            api_key = os.environ.get("ANTHROPIC_API_KEY")
            if api_key and not api_key.startswith("your-") and len(api_key) > 20:
                self.providers["claude"] = {
                    "api_key": api_key,
                    "available": True,
                    "model": os.environ.get("CLAUDE_EMBEDDING_MODEL", "claude-3-embedding-3-0"),
                    "dimension": self.dimensions["claude"],
                    "api_base": os.environ.get("ANTHROPIC_API_BASE", "https://api.anthropic.com/v1")
                }
                print(f"Claude embedding provider ready: {self.providers['claude']['model']} (API key: {api_key[:5]}...{api_key[-4:]})")
            else:
                print(f"Claude embedding provider unavailable: Invalid API key [{api_key[:5] if api_key else 'None'}...]")
                self.providers["claude"] = {"available": False}
        
        # OpenAI API setup
        if "openai" in self.provider_priority:
            # Get API key from .env file - read the actual .env variable
            api_key = os.environ.get("OPENAI_API_KEY")
            if api_key and not api_key.startswith("your-") and len(api_key) > 20:
                self.providers["openai"] = {
                    "api_key": api_key,
                    "available": True,
                    "api_base": os.environ.get("OPENAI_API_BASE", "https://api.openai.com/v1"),
                    "model": os.environ.get("OPENAI_EMBEDDING_MODEL", "text-embedding-ada-002"),
                    "dimension": self.dimensions["openai"]
                }
                print(f"OpenAI embedding provider ready: {self.providers['openai']['model']} (API key: {api_key[:5]}...{api_key[-4:]})")
            else:
                print(f"OpenAI embedding provider unavailable: Invalid API key [{api_key[:5] if api_key else 'None'}...]")
                self.providers["openai"] = {"available": False}

        # Gemini API setup (Google Generative Language API)
        if "gemini" in self.provider_priority:
            gkey = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
            if gkey and len(gkey) > 20:
                self.providers["gemini"] = {
                    "api_key": gkey,
                    "available": True,
                    "api_base": os.environ.get("GEMINI_API_BASE", "https://generativelanguage.googleapis.com/v1beta"),
                    "model": os.environ.get("GEMINI_EMBEDDING_MODEL", "text-embedding-004"),
                    "dimension": self.dimensions["gemini"]
                }
                print("Gemini embedding provider ready: text-embedding-004")
            else:
                print("Gemini embedding provider unavailable: Missing or invalid API key")
                self.providers["gemini"] = {"available": False}
                
    def get_embedding(self, text: str) -> List[float]:
        """
        Get an embedding for a text string using configured provider priority.
        
        Args:
            text: The text to embed
            
        Returns:
            A list of floats representing the embedding vector
        """
        if not text:
            # Return a zero vector if text is empty
            return [0.0] * self.embedding_dimension
        
        # Try each provider in priority order
        for provider in self.provider_priority:
            if provider not in self.providers or not self.providers[provider]["available"]:
                continue
                
            try:
                if provider == "local":
                    embedding = self._get_local_embedding(text)
                    self.embedding_dimension = len(embedding)  # Update dimension based on result
                    return embedding
                elif provider == "claude":
                    embedding = self._get_claude_embedding(text)
                    self.embedding_dimension = len(embedding)  # Update dimension based on result
                    return embedding
                elif provider == "openai":
                    embedding = self._get_openai_embedding(text)
                    self.embedding_dimension = len(embedding)  # Update dimension based on result
                    return embedding
                elif provider == "gemini":
                    embedding = self._get_gemini_embedding(text)
                    self.embedding_dimension = len(embedding)
                    return embedding
            except Exception as e:
                print(f"Error using {provider} embeddings: {str(e)}")
                continue
        
        # Fallback to random if all providers fail
        print("Warning: All embedding providers failed. Using random embeddings.")
        return self._get_random_embedding()
    
    def _get_local_embedding(self, text: str) -> List[float]:
        """Get embedding from local sentence-transformers model."""
        model = self.providers["local"]["model"]
        try:
            embedding = model.encode(text)
        except Exception as e:
            # Force CPU fallback if a CUDA-related error occurs
            if "CUDA" in str(e).upper():
                try:
                    from sentence_transformers import SentenceTransformer
                    model_name = self.model_name
                    self.providers["local"]["model"] = SentenceTransformer(model_name, local_files_only=True, device="cpu")
                    print("Local embedding: CUDA error detected, falling back to CPU")
                    embedding = self.providers["local"]["model"].encode(text)
                except Exception:
                    raise
            else:
                raise
        return embedding.tolist()

    def _get_gemini_embedding(self, text: str) -> List[float]:
        """Get embedding from Google Gemini embeddings API (text-embedding-004)."""
        provider = self.providers.get("gemini", {})
        api_key = provider.get("api_key")
        api_base = provider.get("api_base", "https://generativelanguage.googleapis.com/v1beta")
        model = provider.get("model", "text-embedding-004")
        url = f"{api_base.rstrip('/')}/models/{model}:embedText?key={api_key}"
        payload = {"text": text}
        headers = {"Content-Type": "application/json"}
        resp = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
        if resp.status_code != 200:
            raise Exception(f"Gemini embeddings error {resp.status_code}: {resp.text}")
        data = resp.json()
        # Response shape: { "embedding": { "value": [floats] } }
        if "embedding" in data and isinstance(data["embedding"], dict) and "value" in data["embedding"]:
            return data["embedding"]["value"]
        # Some SDKs return { "embeddings": [ { "value": [...] } ] }
        if "embeddings" in data and isinstance(data["embeddings"], list) and data["embeddings"]:
            emb = data["embeddings"][0]
            if isinstance(emb, dict) and "value" in emb:
                return emb["value"]
        raise Exception(f"Unexpected Gemini embedding response: {data}")

    def _get_claude_embedding(self, text: str) -> List[float]:
        """Get embedding from Claude API."""
        # Claude's API version for embeddings
        # Try the documented approach first (may change over time)
        headers = {
            "Content-Type": "application/json",
            "x-api-key": self.providers["claude"]["api_key"],
            "anthropic-version": "2023-06-01"  # API version may change
        }
        
        # Standard format for the request
        data = {
            "model": self.providers["claude"]["model"],
            "input": text
        }
        
        api_base = self.providers["claude"]["api_base"]
        
        # Try the v1 embeddings endpoint
        embeddings_endpoint = f"{api_base.rstrip('/')}/embeddings"
        
        try:
            print(f"Using Claude embeddings API: {embeddings_endpoint}")
            response = requests.post(
                embeddings_endpoint,
                headers=headers, 
                json=data
            )
            
            if response.status_code == 200:
                result = response.json()
                # Check if the response has the expected structure
                if "embedding" in result:
                    return result["embedding"]
                elif "embeddings" in result and len(result["embeddings"]) > 0:
                    return result["embeddings"][0]
                else:
                    raise Exception(f"Unexpected response format: {result}")
                    
            # Try alternative API path if first attempt fails with 404
            elif response.status_code == 404:
                # Alternative v2 endpoint
                print("Original endpoint not found, trying alternative API version...")
                headers["anthropic-version"] = "2023-01-01"  # Try a different API version
                alt_endpoint = f"{api_base.rstrip('/')}/v1/embeddings"
                
                alt_response = requests.post(
                    alt_endpoint,
                    headers=headers, 
                    json=data
                )
                
                if alt_response.status_code == 200:
                    result = alt_response.json()
                    if "embedding" in result:
                        return result["embedding"]
                    elif "embeddings" in result and len(result["embeddings"]) > 0:
                        return result["embeddings"][0]
                
                # If that also fails, use the Claude completion API to get embeddings
                print("Embeddings API unavailable. Falling back to simulated embedding...")
                return self._get_random_embedding()  # Fall back to random for now
            
            # Other errors
            raise Exception(f"Claude API error: {response.status_code} {response.text}")
        except Exception as e:
            print(f"Claude embedding API error: {str(e)}")
            raise

    def _get_openai_embedding(self, text: str) -> List[float]:
        """Get an embedding from OpenAI API."""
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.providers['openai']['api_key']}"
        }
        
        data = {
            "input": text,
            "model": self.providers["openai"]["model"]
        }
        
        response = requests.post(
            f"{self.providers['openai']['api_base']}/embeddings", 
            headers=headers, 
            json=data
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code} {response.text}")
        
        result = response.json()
        return result["data"][0]["embedding"]
    
    def _get_random_embedding(self) -> List[float]:
        """Generate a random embedding for testing purposes."""
        random_vector = np.random.randn(self.embedding_dimension)
        # Normalize the vector to unit length
        normalized = random_vector / np.linalg.norm(random_vector)
        return normalized.tolist()
    
    def generate_triple_embeddings(self, triple) -> Dict[str, List[float]]:
        """
        Generate embeddings for the subject, predicate, and object of a triple.
        
        Args:
            triple: The Triple object to generate embeddings for
            
        Returns:
            Dictionary with embeddings for subject, predicate, and object
        """
        # Generate embeddings
        subject_embedding = self.get_embedding(triple.subject)
        predicate_embedding = self.get_embedding(triple.predicate)
        
        # Object could be literal or URI
        object_text = triple.object_literal if triple.is_literal else triple.object_uri
        object_embedding = self.get_embedding(object_text)
        
        return {
            "subject_embedding": subject_embedding,
            "predicate_embedding": predicate_embedding,
            "object_embedding": object_embedding
        }
    
    def update_triple_embeddings(self, triple, commit: bool = True):
        """
        Update the embeddings for a triple.
        
        Args:
            triple: The Triple object to update embeddings for
            commit: Whether to commit the session after update
            
        Returns:
            The updated Triple object
        """
        from app import db
        
        embeddings = self.generate_triple_embeddings(triple)
        
        # Update the triple
        triple.subject_embedding = embeddings["subject_embedding"]
        triple.predicate_embedding = embeddings["predicate_embedding"]
        triple.object_embedding = embeddings["object_embedding"]
        
        if commit:
            db.session.commit()
        
        return triple
    
    def batch_update_embeddings(self, triple_ids: List[int] = None, limit: int = 100):
        """
        Update embeddings for multiple triples in batch.
        
        Args:
            triple_ids: Optional list of triple IDs to update. If None, update all triples.
            limit: Maximum number of triples to update at once
            
        Returns:
            Number of triples updated
        """
        from app import db
        from app.models.triple import Triple
        
        query = db.session.query(Triple)
        
        # Filter by IDs if provided
        if triple_ids:
            query = query.filter(Triple.id.in_(triple_ids))
        
        # Filter triples with missing embeddings
        query = query.filter(Triple.subject_embedding.is_(None))
        
        # Limit the batch size
        query = query.limit(limit)
        
        triples = query.all()
        
        print(f"Updating embeddings for {len(triples)} triples...")
        
        for triple in triples:
            self.update_triple_embeddings(triple, commit=False)
        
        db.session.commit()
        
        return len(triples)
    
    def search_similar_chunks(self, query: str, k: int = 5, world_id: Optional[int] = None, document_type: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        Find document chunks similar to the query text.
        
        Args:
            query: The query text to find similar chunks for
            k: Maximum number of results to return
            world_id: Optional world ID to filter chunks by
            document_type: Optional document type to filter by
            
        Returns:
            List of similar document chunks with metadata
        """
        from app import db
        from app.models.document import Document, DocumentChunk
        
        try:
            # Generate embedding for the query text
            embedding = self.get_embedding(query)
            
            # Convert the embedding to a string representation for SQL
            embedding_str = f"[{','.join(str(x) for x in embedding)}]"
            
            # Build the base query
            query_parts = [
                "SELECT dc.id, dc.chunk_text, dc.chunk_index, d.title, d.document_type,",
                "dc.embedding <-> :embedding AS distance",
                "FROM document_chunks dc",
                "JOIN documents d ON dc.document_id = d.id",
                "WHERE dc.embedding IS NOT NULL"
            ]
            
            # Add filters if provided
            params = {"embedding": embedding_str}
            
            if world_id is not None:
                query_parts.append("AND d.world_id = :world_id")
                params["world_id"] = world_id
                
            if document_type is not None:
                query_parts.append("AND d.document_type = :document_type")
                params["document_type"] = document_type
                
            # Add ordering and limit
            query_parts.append("ORDER BY distance")
            query_parts.append("LIMIT :k")
            params["k"] = k
            
            # Combine query parts
            query_str = " ".join(query_parts)
            
            # Execute the query
            result = db.session.execute(text(query_str), params)
            
            # Format results
            similar_chunks = []
            for row in result:
                similar_chunks.append({
                    "id": row.id,
                    "chunk_text": row.chunk_text,
                    "chunk_index": row.chunk_index,
                    "title": row.title,
                    "document_type": row.document_type,
                    "distance": row.distance,
                    "similarity": 1.0 - float(row.distance) if row.distance is not None else 0.0
                })
            
            return similar_chunks
            
        except Exception as e:
            import traceback
            logger.error(f"Error in search_similar_chunks: {str(e)}")
            logger.error(traceback.format_exc())
            return []
    
    def find_similar_triples(self, text: str, field: str = "subject", limit: int = 10) -> List[Dict[str, Any]]:
        """
        Find triples with similar embeddings to the given text.
        
        Args:
            text: The text to find similar triples for
            field: Which field to search (subject, predicate, object)
            limit: Maximum number of results to return
            
        Returns:
            List of (triple, similarity) tuples
        """
        from app import db
        
        # Generate embedding for the query text
        embedding = self.get_embedding(text)
        
        # Determine which embedding field to search
        if field == "subject":
            embedding_field = "subject_embedding"
        elif field == "predicate":
            embedding_field = "predicate_embedding"
        elif field == "object":
            embedding_field = "object_embedding"
        else:
            raise ValueError(f"Invalid field: {field}")
        
        # Convert the embedding to a string representation for SQL
        embedding_str = f"[{','.join(str(x) for x in embedding)}]"
        
        # Query for similar triples
        query = f"""
        SELECT 
            id,
            subject,
            predicate, 
            object_literal,
            object_uri,
            is_literal,
            {embedding_field} <-> '{embedding_str}'::vector AS distance
        FROM 
            character_triples
        WHERE 
            {embedding_field} IS NOT NULL
        ORDER BY 
            distance
        LIMIT {limit}
        """
        
        result = db.session.execute(text(query))
        
        # Format results
        similar_triples = []
        for row in result:
            object_value = row.object_literal if row.is_literal else row.object_uri
            similar_triples.append({
                "id": row.id,
                "subject": row.subject,
                "predicate": row.predicate,
                "object": object_value,
                "is_literal": row.is_literal,
                "similarity": 1.0 - row.distance
            })
        
        return similar_triples
    
    def _split_text(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """
        Split text into smaller chunks for processing.
        
        Args:
            text: The text to split
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        # Simple paragraph-based chunking
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # If adding this paragraph would exceed chunk size, store current chunk and start a new one
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Include overlap from the end of previous chunk
                if len(current_chunk) > chunk_overlap:
                    current_chunk = current_chunk[-chunk_overlap:] + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def embed_documents(self, chunks: List[str]) -> List[List[float]]:
        """
        Generate embeddings for a list of text chunks.
        
        Args:
            chunks: List of text chunks to embed
            
        Returns:
            List of embedding vectors
        """
        embeddings = []
        
        print(f"Generating embeddings for {len(chunks)} chunks...")
        
        for chunk in chunks:
            try:
                embedding = self.get_embedding(chunk)
                embeddings.append(embedding)
            except Exception as e:
                print(f"Error generating embedding for chunk: {str(e)}")
                # Use zero vector as fallback for failed embeddings
                embeddings.append([0.0] * self.embedding_dimension)
        
        return embeddings
    
    def _store_chunks(self, document_id: int, chunks: List[str], embeddings: List[List[float]]) -> int:
        """
        Store document chunks with their embeddings in the database.
        
        Args:
            document_id: ID of the document these chunks belong to
            chunks: List of text chunks
            embeddings: List of embedding vectors corresponding to chunks
            
        Returns:
            Number of chunks stored
        """
        from app import db
        from app.models.document import DocumentChunk
        
        # Make sure we have the same number of chunks and embeddings
        if len(chunks) != len(embeddings):
            raise ValueError(f"Number of chunks ({len(chunks)}) does not match number of embeddings ({len(embeddings)})")
        
        # Delete existing chunks for this document
        existing_chunks = DocumentChunk.query.filter_by(document_id=document_id).all()
        for chunk in existing_chunks:
            db.session.delete(chunk)
        
        # Store new chunks
        for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                document_id=document_id,
                chunk_index=i,
                content=chunk_text,
                embedding=embedding,
                chunk_metadata={"position": i}
            )
            db.session.add(chunk)
        
        db.session.commit()
        return len(chunks)
    
    def process_url(self, url: str, title: str, document_type: str, world_id: int) -> int:
        """
        Process a URL into a document with embeddings.
        
        Args:
            url: The URL to process
            title: Title for the document
            document_type: Type of document (guideline, case_study, etc.)
            world_id: ID of the world this document belongs to
            
        Returns:
            ID of the created document
        """
        from app import db
        from app.models.document import Document, PROCESSING_STATUS
        
        # Create document record
        document = Document(
            title=title,
            document_type=document_type,
            world_id=world_id,
            source=url,
            file_type="url",
            doc_metadata={},
            processing_status=PROCESSING_STATUS['PROCESSING']
        )
        db.session.add(document)
        db.session.commit()
        
        try:
            # Extract text from URL
            text = self._extract_from_url(url)
            document.content = text
            
            # Split text into chunks
            chunks = self._split_text(text)
            
            # Generate embeddings
            embeddings = self.embed_documents(chunks)
            
            # Store chunks with embeddings
            self._store_chunks(document.id, chunks, embeddings)
            
            # Update document status
            document.processing_status = PROCESSING_STATUS['COMPLETED']
            document.processing_progress = 100
            db.session.commit()
            
            return document.id
            
        except Exception as e:
            # Update document status to failed
            document.processing_status = PROCESSING_STATUS['FAILED']
            document.processing_error = str(e)
            db.session.commit()
            
            # Re-raise the exception
            raise
